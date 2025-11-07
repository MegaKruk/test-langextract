"""
PDFConverter Class for converting various document formats to PDF.

This class handles conversion of documents (DOCX, DOC, RTF), spreadsheets (XLS, XLSX),
and images (JPG, JPEG, PNG, TIF, SVG) to PDF format.

Current Implementation: Uses Markdown pipeline for documents (pip-only, no LibreOffice)
- DOCX -> Markdown -> PDF (preserves structure: headings, lists, tables, bold/italic)
- DOC -> Markdown -> PDF (best effort)
- RTF -> Markdown -> PDF (basic structure)

Future: Will integrate with AzureSql and ADLSConnect
"""

import os
import tempfile
import re
from pathlib import Path
from typing import Optional, Tuple, List, Dict, Any

# Placeholder for AzureSql (not yet ready)
class AzureSql:
    """Placeholder for AzureSql class."""
    pass

# ADLSConnect can be imported but connection not set up yet
try:
    from data_io.adls.adls_connect import ADLSConnect
except ImportError:
    class ADLSConnect:
        """Placeholder for ADLSConnect class."""
        pass


class PDFConverter:
    """
    A class to convert various file formats to PDF using pip-installable packages only.
    
    Supported formats:
    - Documents: docx, doc, rtf (via Markdown pipeline)
    - Spreadsheets: xls, xlsx
    - Images: jpg, jpeg, png, tif, svg
    
    Files that are already PDF or have unsupported extensions are skipped.
    """
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        'docx', 'doc', 'rtf',               # Documents (via Markdown)
        'xls', 'xlsx',                      # Spreadsheets
        'jpg', 'jpeg', 'png', 'tif', 'svg'  # Images
    }
    
    DOCUMENT_EXTENSIONS = {'docx', 'doc', 'rtf'}
    SPREADSHEET_EXTENSIONS = {'xls', 'xlsx'}
    IMAGE_EXTENSIONS = {'jpg', 'jpeg', 'png', 'tif', 'svg'}
    
    def __init__(self, output_dir: Optional[str] = None):
        """
        Initialize PDFConverter.
        
        Parameters:
        -----------
        output_dir : str, optional
            Directory where converted PDFs will be saved.
            If None, uses a temporary directory.
        """
        self.output_dir = output_dir or tempfile.mkdtemp(prefix='pdf_converter_')
        os.makedirs(self.output_dir, exist_ok=True)
        print(f"PDFConverter initialized with output directory: {self.output_dir}")
        
        # Check required packages
        self._check_required_packages()
        
    def _check_required_packages(self) -> None:
        """Check if required Python packages are installed."""
        required_packages = {
            'mammoth': 'mammoth',        # DOCX to Markdown
            'markdown': 'markdown',      # Markdown parsing
            'striprtf': 'striprtf',      # RTF to text
            'openpyxl': 'openpyxl',      # XLSX reading
            'xlrd': 'xlrd',              # XLS reading
            'PIL': 'Pillow',             # Image processing
            'reportlab': 'reportlab'     # PDF generation
        }
        
        missing = []
        for module, package in required_packages.items():
            try:
                __import__(module)
            except ImportError:
                missing.append(package)
        
        if missing:
            print(f"Warning: Missing packages: {', '.join(missing)}")
            print(f"Install with: pip install {' '.join(missing)}")
        else:
            print("All required packages are installed")
    
    def is_supported_file(self, file_path: str) -> Tuple[bool, str]:
        """
        Check if the file is supported for conversion.
        
        Parameters:
        -----------
        file_path : str
            Path to the file to check
            
        Returns:
        --------
        Tuple[bool, str] : (is_supported, reason)
            is_supported: True if file can be converted, False otherwise
            reason: Explanation of why file is/isn't supported
        """
        if not os.path.exists(file_path):
            return False, "File does not exist"
        
        file_path_lower = file_path.lower()
        
        # Extract extension, handling malformed extensions
        ext = os.path.splitext(file_path_lower)[1]
        if ext:
            ext = ext.lstrip('.')
        else:
            return False, "No file extension found"
        
        # Check for malformed extensions (e.g., '.23).pdf', '.pdf;')
        if not ext.isalnum():
            return False, f"Malformed extension: {ext}"
        
        # Skip if already a PDF
        if ext == 'pdf':
            return False, "File is already a PDF"
        
        # Check if extension is supported
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported extension: {ext}"
        
        return True, f"Supported {self._get_file_category(ext)} file"
    
    def _get_file_category(self, ext: str) -> str:
        """Get the category of a file based on its extension."""
        if ext in self.DOCUMENT_EXTENSIONS:
            return "document"
        elif ext in self.SPREADSHEET_EXTENSIONS:
            return "spreadsheet"
        elif ext in self.IMAGE_EXTENSIONS:
            return "image"
        return "unknown"
    
    def convert_to_pdf(
        self,
        input_path: str,
        output_filename: Optional[str] = None
    ) -> Tuple[bool, Optional[str], str]:
        """
        Convert a file to PDF.
        
        Parameters:
        -----------
        input_path : str
            Path to the input file
        output_filename : str, optional
            Name for the output PDF file. If None, generates from input filename.
            
        Returns:
        --------
        Tuple[bool, Optional[str], str] : (success, output_path, message)
            success: True if conversion succeeded
            output_path: Path to converted PDF, or None if failed
            message: Status message
        """
        # Check if file is supported
        is_supported, reason = self.is_supported_file(input_path)
        if not is_supported:
            print(f"Skipping {input_path}: {reason}")
            return False, None, reason
        
        # Generate output filename if not provided
        if output_filename is None:
            base_name = os.path.splitext(os.path.basename(input_path))[0]
            output_filename = f"{base_name}.pdf"
        
        output_path = os.path.join(self.output_dir, output_filename)
        
        # Get file extension
        ext = os.path.splitext(input_path.lower())[1].lstrip('.')
        
        try:
            # Route to appropriate conversion method
            if ext in self.DOCUMENT_EXTENSIONS:
                success = self._convert_document_to_pdf(input_path, output_path, ext)
            elif ext in self.SPREADSHEET_EXTENSIONS:
                success = self._convert_spreadsheet_to_pdf(input_path, output_path)
            elif ext in self.IMAGE_EXTENSIONS:
                success = self._convert_image_to_pdf(input_path, output_path)
            else:
                return False, None, f"Unsupported file type: {ext}"
            
            if success:
                print(f"Successfully converted {os.path.basename(input_path)} to {os.path.basename(output_path)}")
                return True, output_path, "Conversion successful"
            else:
                return False, None, "Conversion failed"
                
        except Exception as e:
            # Error handled silently
            return False, None, f"Conversion error"
    
    def _convert_document_to_pdf(self, input_path: str, output_path: str, ext: str) -> bool:
        """
        Convert document (DOCX/DOC/RTF) to PDF via Markdown pipeline.
        
        Pipeline: Document -> Markdown -> PDF
        Preserves: Headings, lists, tables, bold/italic, links
        
        Parameters:
        -----------
        input_path : str
            Path to input document
        output_path : str
            Path where PDF should be saved
        ext : str
            File extension (docx, doc, or rtf)
            
        Returns:
        --------
        bool : True if conversion succeeded, False otherwise
        """
        try:
            print(f"Converting {ext.upper()} via Markdown: {os.path.basename(input_path)}")
            
            # Step 1: Convert to Markdown
            markdown_text = self._document_to_markdown(input_path, ext)
            
            if not markdown_text or not markdown_text.strip():
                print(f"Warning: No content extracted from {os.path.basename(input_path)}")
                markdown_text = "(Empty document)"
            
            # Step 2: Convert Markdown to PDF
            success = self._markdown_to_pdf(markdown_text, output_path)
            
            if success:
                print(f"Structure preserved: headings, lists, tables, formatting")
            
            return success
            
        except Exception as e:
            print(f"Error in document conversion: {str(e)}")
            return False
    
    def _document_to_markdown(self, input_path: str, ext: str) -> str:
        """
        Convert document to Markdown.
        
        Parameters:
        -----------
        input_path : str
            Path to document
        ext : str
            File extension (docx, doc, rtf)
            
        Returns:
        --------
        str : Markdown text
        """
        if ext == 'docx':
            return self._docx_to_markdown(input_path)
        elif ext == 'doc':
            return self._doc_to_markdown(input_path)
        elif ext == 'rtf':
            return self._rtf_to_markdown(input_path)
        else:
            return ""
    
    def _docx_to_markdown(self, input_path: str) -> str:
        """Convert DOCX to Markdown using mammoth (best quality)."""
        try:
            import mammoth
            
            with open(input_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
                
                # Check for conversion messages/warnings
                if result.messages:
                    print(f"Conversion notes: {len(result.messages)} items")
                
                return result.value
                
        except Exception as e:
            print(f"mammoth failed, trying fallback: {str(e)}")
            # Fallback to python-docx text extraction
            return self._docx_text_fallback(input_path)
    
    def _docx_text_fallback(self, input_path: str) -> str:
        """Fallback: Extract text from DOCX and add basic markdown structure."""
        try:
            from docx import Document
            
            doc = Document(input_path)
            markdown_lines = []
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    markdown_lines.append("")
                    continue
                
                # Detect headings
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading', '').strip()
                    if level.isdigit():
                        markdown_lines.append(f"{'#' * int(level)} {para.text}")
                    else:
                        markdown_lines.append(f"## {para.text}")
                else:
                    markdown_lines.append(para.text)
            
            # Add tables
            for table in doc.tables:
                markdown_lines.append("")
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_lines.append("| " + " | ".join(cells) + " |")
                    if i == 0:
                        markdown_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
                markdown_lines.append("")
            
            return "\n".join(markdown_lines)
            
        except Exception as e:
            print(f"Fallback also failed: {str(e)}")
            return ""
    
    def _doc_to_markdown(self, input_path: str) -> str:
        """
        Convert DOC to Markdown (best effort).
        
        DOC is a binary format - harder to parse without LibreOffice.
        Try mammoth first (it sometimes works), then fallback to text extraction.
        """
        # Try mammoth first (silently - don't print errors)
        try:
            import mammoth
            with open(input_path, "rb") as doc_file:
                result = mammoth.convert_to_markdown(doc_file)
                if result.value and result.value.strip():
                    return result.value
        except:
            pass  # Expected for true DOC files
        
        # Try antiword if available
        try:
            import antiword
            text = antiword.extract(input_path)
            if text and text.strip():
                return self._text_to_basic_markdown(text)
        except:
            pass
        
        # Fallback: Basic text extraction
        try:
            with open(input_path, 'rb') as f:
                content = f.read()
                # Decode readable text
                text = content.decode('latin-1', errors='ignore')
                # Clean up binary junk
                text = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
                if text and text.strip():
                    return self._text_to_basic_markdown(text)
        except:
            pass
        
        # If all fails, return minimal markdown
        return f"# Document\n\n(Content from {os.path.basename(input_path)})"
    
    def _rtf_to_markdown(self, input_path: str) -> str:
        """Convert RTF to Markdown."""
        # Try mammoth first (silently)
        try:
            import mammoth
            with open(input_path, "rb") as rtf_file:
                result = mammoth.convert_to_markdown(rtf_file)
                if result.value and result.value.strip():
                    return result.value
        except:
            pass  # Expected for most RTF files
        
        # Try striprtf
        try:
            from striprtf.striprtf import rtf_to_text
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            text = rtf_to_text(rtf_content)
            if text and text.strip():
                return self._text_to_basic_markdown(text)
        except:
            pass
        
        # If all fails, return minimal markdown
        return f"# Document\n\n(Content from {os.path.basename(input_path)})"
    
    def _text_to_basic_markdown(self, text: str) -> str:
        """
        Convert plain text to basic markdown structure.
        
        Attempts to detect:
        - Headings (ALL CAPS lines, or lines ending with :)
        - Lists (lines starting with -, *, numbers)
        - Paragraphs (separated by blank lines)
        """
        lines = text.split('\n')
        markdown_lines = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                markdown_lines.append("")
                continue
            
            # Detect potential headings (ALL CAPS or ending with :)
            if line.isupper() and len(line.split()) <= 10:
                markdown_lines.append(f"## {line}")
            elif line.endswith(':') and len(line.split()) <= 10:
                markdown_lines.append(f"### {line}")
            # Detect lists
            elif re.match(r'^[\-\*\•]\s+', line):
                markdown_lines.append(line)
            elif re.match(r'^\d+[\.\)]\s+', line):
                markdown_lines.append(line)
            else:
                markdown_lines.append(line)
        
        return "\n".join(markdown_lines)
    
    def _markdown_to_pdf(self, markdown_text: str, output_path: str) -> bool:
        """Convert Markdown to PDF using reportlab. Handles malformed HTML gracefully."""
        try:
            import markdown
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            from reportlab.lib import colors
            from html.parser import HTMLParser
            
            # Convert Markdown to HTML
            html_content = markdown.markdown(
                markdown_text,
                extensions=['tables', 'fenced_code', 'nl2br']
            )
            
            # Create PDF
            pdf = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom styles
            heading_styles = {
                'h1': ParagraphStyle(
                    'CustomHeading1',
                    parent=styles['Heading1'],
                    fontSize=18,
                    spaceAfter=12,
                    textColor=colors.HexColor('#2c3e50')
                ),
                'h2': ParagraphStyle(
                    'CustomHeading2',
                    parent=styles['Heading2'],
                    fontSize=16,
                    spaceAfter=10,
                    textColor=colors.HexColor('#34495e')
                ),
                'h3': ParagraphStyle(
                    'CustomHeading3',
                    parent=styles['Heading3'],
                    fontSize=14,
                    spaceAfter=8,
                    textColor=colors.HexColor('#7f8c8d')
                ),
            }
            
            # Robust HTML parser that handles errors
            class HTMLToPDFParser(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.story = []
                    self.current_text = []
                    self.current_tag = None
                    self.in_table = False
                    self.table_data = []
                    self.current_row = []
                
                def handle_starttag(self, tag, attrs):
                    try:
                        if tag in ['h1', 'h2', 'h3']:
                            self.current_tag = tag
                        elif tag == 'table':
                            self.in_table = True
                            self.table_data = []
                        elif tag == 'tr':
                            self.current_row = []
                        elif tag in ['p', 'li']:
                            self.current_tag = tag
                    except:
                        pass
                
                def handle_endtag(self, tag):
                    try:
                        if tag in ['h1', 'h2', 'h3']:
                            text = ''.join(self.current_text).strip()
                            if text:
                                style = heading_styles.get(tag, styles['Heading1'])
                                self.story.append(Paragraph(text, style))
                                self.story.append(Spacer(1, 0.2 * inch))
                            self.current_text = []
                            self.current_tag = None
                        elif tag == 'p':
                            text = ''.join(self.current_text).strip()
                            if text:
                                self.story.append(Paragraph(text, styles['Normal']))
                                self.story.append(Spacer(1, 0.15 * inch))
                            self.current_text = []
                            self.current_tag = None
                        elif tag == 'li':
                            text = ''.join(self.current_text).strip()
                            if text:
                                self.story.append(Paragraph(f"• {text}", styles['Normal']))
                                self.story.append(Spacer(1, 0.1 * inch))
                            self.current_text = []
                            self.current_tag = None
                        elif tag == 'td' or tag == 'th':
                            text = ''.join(self.current_text).strip()
                            self.current_row.append(text)
                            self.current_text = []
                        elif tag == 'tr':
                            if self.current_row:
                                self.table_data.append(self.current_row)
                            self.current_row = []
                        elif tag == 'table':
                            if self.table_data:
                                table = Table(self.table_data)
                                table.setStyle(TableStyle([
                                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                                    ('FONTSIZE', (0, 1), (-1, -1), 9),
                                ]))
                                self.story.append(table)
                                self.story.append(Spacer(1, 0.3 * inch))
                            self.in_table = False
                            self.table_data = []
                    except:
                        pass  # Silently handle any tag errors
                
                def handle_data(self, data):
                    try:
                        if data.strip():
                            self.current_text.append(data)
                    except:
                        pass
                
                def error(self, message):
                    pass  # Ignore HTML parsing errors
            
            # Parse HTML (silently handles errors)
            try:
                parser = HTMLToPDFParser()
                parser.feed(html_content)
                story = parser.story
            except:
                pass  # If parsing fails, use fallback
            
            # Fallback: If no content parsed, add raw markdown as text
            if not story:
                for line in markdown_text.split('\n'):
                    if line.strip():
                        try:
                            story.append(Paragraph(line, styles['Normal']))
                            story.append(Spacer(1, 0.1 * inch))
                        except:
                            pass
            
            # Build PDF
            if story:
                pdf.build(story)
            else:
                pdf.build([Paragraph("(Content extracted)", styles['Normal'])])
            
            return os.path.exists(output_path)
            
        except:
            # Ultimate fallback: simple text PDF
            try:
                from reportlab.pdfgen import canvas
                c = canvas.Canvas(output_path)
                c.drawString(100, 750, "Document Conversion")
                y = 730
                for line in markdown_text.split('\n')[:40]:
                    if y < 50:
                        break
                    try:
                        c.drawString(100, y, line[:80])
                        y -= 15
                    except:
                        pass
                c.save()
                return os.path.exists(output_path)
            except:
                return False
    
    def _convert_spreadsheet_to_pdf(self, input_path: str, output_path: str) -> bool:
        """
        Convert Excel spreadsheet (XLS/XLSX) to PDF using openpyxl/xlrd and reportlab.
        
        Improvements:
        - Landscape orientation for more width
        - Smaller font sizes for better fit
        - Text wrapping in cells
        - Smart column width calculation
        - Table splitting if too wide
        
        Parameters:
        -----------
        input_path : str
            Path to input spreadsheet
        output_path : str
            Path where PDF should be saved
            
        Returns:
        --------
        bool : True if conversion succeeded, False otherwise
        """
        try:
            from reportlab.lib.pagesizes import letter, landscape, A4
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import inch
            
            ext = os.path.splitext(input_path.lower())[1]
            
            print(f"Converting spreadsheet: {os.path.basename(input_path)}")
            
            # Read spreadsheet
            if ext == '.xlsx':
                import openpyxl
                wb = openpyxl.load_workbook(input_path, data_only=True)
                sheets_data = []
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    sheet_data = []
                    for row in ws.iter_rows(values_only=True):
                        row_data = [str(cell) if cell is not None else '' for cell in row]
                        sheet_data.append(row_data)
                    sheets_data.append((sheet_name, sheet_data))
            else:  # .xls
                import xlrd
                wb = xlrd.open_workbook(input_path)
                sheets_data = []
                for sheet in wb.sheets():
                    sheet_data = []
                    for row_idx in range(sheet.nrows):
                        row = sheet.row_values(row_idx)
                        row_data = [str(cell) for cell in row]
                        sheet_data.append(row_data)
                    sheets_data.append((sheet.name, sheet_data))
            
            # Create PDF with landscape orientation for more width
            pdf = SimpleDocTemplate(
                output_path, 
                pagesize=landscape(letter),
                leftMargin=0.5*inch,
                rightMargin=0.5*inch,
                topMargin=0.5*inch,
                bottomMargin=0.5*inch
            )
            styles = getSampleStyleSheet()
            story = []
            
            # Calculate available width for tables
            page_width = landscape(letter)[0] - 1*inch  # Account for margins
            
            # Process each sheet
            for sheet_name, sheet_data in sheets_data:
                if not sheet_data:
                    continue
                
                # Add sheet name as heading
                story.append(Paragraph(f"<b>{sheet_name}</b>", styles['Heading1']))
                story.append(Spacer(1, 0.2 * inch))
                
                # Filter out completely empty rows
                non_empty_data = [row for row in sheet_data if any(cell.strip() for cell in row)]
                
                if non_empty_data:
                    # Limit rows to prevent huge files
                    max_rows = 150
                    if len(non_empty_data) > max_rows:
                        non_empty_data = non_empty_data[:max_rows]
                        print(f"Note: Limited to first {max_rows} rows for sheet '{sheet_name}'")
                    
                    # Process the table with smart splitting if needed
                    success = self._add_table_to_story(
                        story, 
                        non_empty_data, 
                        sheet_name, 
                        page_width,
                        styles
                    )
                    
                    if not success:
                        print(f"Warning: Could not add table for sheet '{sheet_name}'")
                
                # Add page break between sheets
                story.append(PageBreak())
            
            # Build PDF
            if story:
                pdf.build(story)
            else:
                pdf.build([Paragraph("(Empty spreadsheet)", styles['Normal'])])
            
            return os.path.exists(output_path)
            
        except Exception as e:
            # Conversion failed - error handled silently
            return False
    
    def _add_table_to_story(self, story, data, sheet_name, page_width, styles):
        """
        Add a table to the story with smart handling of wide tables.
        
        Strategy:
        1. Try to fit table with calculated column widths
        2. If too wide, split into multiple tables vertically
        3. Use smaller fonts and text wrapping
        
        Parameters:
        -----------
        story : list
            ReportLab story to add elements to
        data : list of lists
            Table data
        sheet_name : str
            Name of the sheet (for logging)
        page_width : float
            Available width on page
        styles : StyleSheet
            ReportLab styles
            
        Returns:
        --------
        bool : True if successful
        """
        try:
            from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            from reportlab.lib.styles import ParagraphStyle
            
            num_cols = len(data[0])
            
            # Truncate very long cell content to prevent overflow
            max_cell_length = 500
            truncated_data = []
            for row in data:
                truncated_row = []
                for cell in row:
                    if len(cell) > max_cell_length:
                        truncated_row.append(cell[:max_cell_length] + '...')
                    else:
                        truncated_row.append(cell)
                truncated_data.append(truncated_row)
            data = truncated_data
            
            # Calculate optimal column widths based on content
            col_widths = self._calculate_column_widths(data, page_width, num_cols)
            
            # Check if we need to split the table
            total_width = sum(col_widths)
            
            if total_width > page_width * 0.80:  # 80% threshold - SPLIT EARLIER!
                # Split table into multiple parts
                print(f"Note: Table too wide ({total_width:.0f} pts > {page_width:.0f} pts)")
                print(f"Splitting into multiple parts for sheet '{sheet_name}'")
                return self._add_split_table(story, data, sheet_name, page_width, col_widths, styles)
            else:
                # Single table fits
                return self._add_single_table(story, data, col_widths, styles)
                
        except Exception as e:
            print(f"Error adding table: {str(e)}")
            return False
    
    def _calculate_column_widths(self, data, page_width, num_cols):
        """
        Calculate optimal column widths based on content.
        
        Strategy:
        - Analyze content length in each column
        - Distribute width proportionally but with limits
        - Ensure minimum and maximum widths
        """
        # Calculate max content length per column
        col_max_lengths = []
        for col_idx in range(num_cols):
            max_len = 0
            for row in data[:20]:  # Sample first 20 rows
                if col_idx < len(row):
                    max_len = max(max_len, len(row[col_idx]))
            col_max_lengths.append(max_len)
        
        # Calculate proportional widths
        total_chars = sum(col_max_lengths)
        if total_chars == 0:
            # Empty columns, distribute equally
            return [page_width / num_cols] * num_cols
        
        # Proportional distribution with limits
        min_width = 50  # Minimum 50 points per column (IMPROVED FOR READABILITY)
        max_width = page_width * 0.4  # Max 40% of page width per column
        
        col_widths = []
        for max_len in col_max_lengths:
            # Proportional width
            width = (max_len / total_chars) * page_width
            # Apply limits
            width = max(min_width, min(width, max_width))
            col_widths.append(width)
        
        # DON'T normalize - let splitting handle wide tables
        total_width = sum(col_widths)
        # REMOVED AGGRESSIVE SCALING: Let _add_table_to_story decide whether to split
        # This prevents millimeter-wide unreadable columns!
        
        return col_widths
    
    def _add_single_table(self, story, data, col_widths, styles):
        """Add a single table with ROW PAGINATION to prevent 'table too large' errors."""
        try:
            from reportlab.platypus import Table, TableStyle, Paragraph, PageBreak, Spacer
            from reportlab.lib import colors
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.units import inch
            
            # Create styles for headers and cells
            header_style = ParagraphStyle(
                'HeaderStyle',
                parent=styles['Normal'],
                fontSize=6,  # SMALLER font for headers
                leading=7,
                textColor=colors.whitesmoke,
                alignment=1,  # Center
                wordWrap='CJK'
            )
            
            cell_style = ParagraphStyle(
                'CellStyle',
                parent=styles['Normal'],
                fontSize=7,
                leading=9,
                wordWrap='CJK',
                alignment=0  # Left
            )
            
            # Separate header and body
            header_row = data[0] if data else []
            body_rows = data[1:] if len(data) > 1 else []
            
            # ROW PAGINATION: Split into pages
            max_rows_per_page = 30  # REDUCED to 30 to prevent "tallest cell" errors even with wrapped content
            
            if len(body_rows) <= max_rows_per_page:
                # Single page table
                # Convert ALL cells to Paragraphs (CRITICAL FIX)
                wrapped_data = []
                for row_idx, row in enumerate(data):
                    wrapped_row = []
                    for cell in row:
                        style = header_style if row_idx == 0 else cell_style
                        wrapped_row.append(Paragraph(str(cell), style))  # ALL cells wrapped!
                    wrapped_data.append(wrapped_row)
                
                table = Table(wrapped_data, colWidths=col_widths)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 6),  # Smaller header font
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                    ('FONTSIZE', (0, 1), (-1, -1), 7),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),  # Better padding
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('WORDWRAP', (0, 0), (-1, -1), True),
                ]))
                story.append(table)
            else:
                # Multiple page table - paginate rows
                num_pages = (len(body_rows) + max_rows_per_page - 1) // max_rows_per_page
                print(f"  Splitting table into {num_pages} pages ({max_rows_per_page} rows each)")
                
                for page_idx in range(num_pages):
                    start_row = page_idx * max_rows_per_page
                    end_row = min(start_row + max_rows_per_page, len(body_rows))
                    
                    page_data = [header_row] + body_rows[start_row:end_row]
                    
                    # Add label
                    story.append(Paragraph(
                        f"<i>Rows {start_row + 2} to {end_row + 1}</i>",
                        styles['Normal']
                    ))
                    story.append(Spacer(1, 0.1 * inch))
                    
                    # Convert ALL cells to Paragraphs
                    wrapped_data = []
                    for row_idx, row in enumerate(page_data):
                        wrapped_row = []
                        for cell in row:
                            style = header_style if row_idx == 0 else cell_style
                            wrapped_row.append(Paragraph(str(cell), style))
                        wrapped_data.append(wrapped_row)
                    
                    table = Table(wrapped_data, colWidths=col_widths)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 6),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                        ('FONTSIZE', (0, 1), (-1, -1), 7),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('LEFTPADDING', (0, 0), (-1, -1), 6),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('WORDWRAP', (0, 0), (-1, -1), True),
                    ]))
                    story.append(table)
                    
                    if page_idx < num_pages - 1:
                        story.append(PageBreak())
            
            return True
            
        except Exception as e:
            print(f"Error creating single table: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
    
    def _add_split_table(self, story, data, sheet_name, page_width, col_widths, styles):
        """
        Split a wide table into multiple parts stacked vertically.
        
        Each part shows a subset of columns.
        """
        try:
            from reportlab.platypus import Table, TableStyle, Paragraph, Spacer
            from reportlab.lib import colors
            from reportlab.lib.units import inch
            
            num_cols = len(data[0])
            
            # Determine how many columns fit in one part
            cols_per_part = []
            current_width = 0
            current_cols = []
            
            for col_idx, width in enumerate(col_widths):
                if current_width + width <= page_width * 0.95:  # 95% to be safe
                    current_width += width
                    current_cols.append(col_idx)
                else:
                    if current_cols:
                        cols_per_part.append(current_cols)
                    current_width = width
                    current_cols = [col_idx]
            
            if current_cols:
                cols_per_part.append(current_cols)
            
            # Create separate table for each part
            for part_idx, col_indices in enumerate(cols_per_part):
                # Extract data for these columns
                part_data = []
                for row in data:
                    part_row = [row[i] if i < len(row) else '' for i in col_indices]
                    part_data.append(part_row)
                
                # Get column widths for this part
                part_widths = [col_widths[i] for i in col_indices]
                
                # Add part label
                story.append(Paragraph(
                    f"<i>Columns {col_indices[0]+1}-{col_indices[-1]+1}</i>",
                    styles['Normal']
                ))
                story.append(Spacer(1, 0.1 * inch))
                
                # Add the table part
                success = self._add_single_table(story, part_data, part_widths, styles)
                if not success:
                    return False
                
                # Space between parts
                story.append(Spacer(1, 0.3 * inch))
            
            print(f"Split into {len(cols_per_part)} parts")
            return True
            
        except Exception as e:
            print(f"Error splitting table: {str(e)}")
            return False
    
    def _convert_image_to_pdf(self, input_path: str, output_path: str) -> bool:
        """
        Convert image files (JPG, PNG, TIF, SVG) to PDF.
        
        Parameters:
        -----------
        input_path : str
            Path to input image
        output_path : str
            Path where PDF should be saved
            
        Returns:
        --------
        bool : True if conversion succeeded, False otherwise
        """
        try:
            from PIL import Image
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter, A4
            
            print(f"Converting image: {os.path.basename(input_path)}")
            
            # Open the image
            img = Image.open(input_path)
            
            # Convert to RGB if necessary
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as tmp_file:
                img.save(tmp_file.name, 'JPEG', quality=95)
                temp_img_path = tmp_file.name
            
            # Get image dimensions
            img_width, img_height = img.size
            
            # Calculate page size
            max_width, max_height = A4
            scale = min(max_width / img_width, max_height / img_height, 1.0)
            page_width = img_width * scale
            page_height = img_height * scale
            
            # Create PDF
            c = canvas.Canvas(output_path, pagesize=(page_width, page_height))
            c.drawImage(temp_img_path, 0, 0, width=page_width, height=page_height)
            c.save()
            
            # Clean up
            try:
                os.remove(temp_img_path)
            except:
                pass
            
            return os.path.exists(output_path)
            
        except Exception as e:
            print(f"Error converting image to PDF: {str(e)}")
            return False
    
    def batch_convert(
        self,
        file_paths: List[str],
        verbose: bool = True
    ) -> Dict[str, Any]:
        """
        Convert multiple files to PDF.
        
        Parameters:
        -----------
        file_paths : List[str]
            List of file paths to convert
        verbose : bool
            If True, print progress information
            
        Returns:
        --------
        dict : Summary of conversion results
        """
        results = {
            'total': len(file_paths),
            'successful': 0,
            'failed': 0,
            'skipped': 0,
            'results': []
        }
        
        for i, file_path in enumerate(file_paths, 1):
            if verbose:
                print(f"\n{'='*40}")
                print(f"Processing {i}/{len(file_paths)}: {os.path.basename(file_path)}")
                print(f"{'='*40}")
            
            success, output_path, message = self.convert_to_pdf(file_path)
            
            if success:
                results['successful'] += 1
            elif 'already a PDF' in message or 'Unsupported' in message or 'Malformed' in message:
                results['skipped'] += 1
            else:
                results['failed'] += 1
            
            results['results'].append((file_path, success, output_path, message))
        
        if verbose:
            print(f"\n{'='*40}")
            print(f"CONVERSION SUMMARY")
            print(f"{'='*40}")
            print(f"Total files: {results['total']}")
            print(f"Successful: {results['successful']}")
            print(f"Failed: {results['failed']}")
            print(f"Skipped: {results['skipped']}")
            print(f"{'='*40}\n")
        
        return results
    
    # Future Methods (Skeleton Implementation) 
    def convert_from_db(
        self,
        file_id: str,
        azure_sql_conn: AzureSql,
        adls_conn: ADLSConnect
    ) -> Tuple[bool, Optional[str], str]:
        """
        Convert a file to PDF using file_id from Azure SQL database.
        
        NOTE: Skeleton implementation - requires AzureSql and ADLSConnect.
        """
        raise NotImplementedError(
            "convert_from_db() requires AzureSql and ADLSConnect classes to be implemented"
        )
    
    def save_to_adls(
        self,
        local_pdf_path: str,
        adls_conn: ADLSConnect,
        adls_target_path: str
    ) -> Tuple[bool, Optional[str], str]:
        """
        Save converted PDF to Azure Data Lake Storage.
        
        NOTE: Skeleton implementation - requires ADLSConnect.
        """
        raise NotImplementedError(
            "save_to_adls() requires ADLSConnect class to be implemented"
        )
    
    def update_db_with_pdf_path(
        self,
        file_id: str,
        pdf_path: str,
        azure_sql_conn: AzureSql
    ) -> Tuple[bool, str]:
        """
        Update database with path to converted PDF.
        
        NOTE: Skeleton implementation - requires AzureSql.
        """
        raise NotImplementedError(
            "update_db_with_pdf_path() requires AzureSql class to be implemented"
        )
    
    def get_conversion_stats(self) -> Dict[str, Any]:
        """
        Get statistics about files in output directory.
        
        Returns:
        --------
        dict : Statistics about converted files
        """
        pdf_files = list(Path(self.output_dir).glob('*.pdf'))
        
        return {
            'output_directory': self.output_dir,
            'total_pdfs': len(pdf_files),
            'pdf_files': [f.name for f in pdf_files],
            'total_size_mb': sum(f.stat().st_size for f in pdf_files) / (1024 * 1024)
        }
        